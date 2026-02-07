"""
DOCX Parser for Resume Intelligence Engine
"""
from pathlib import Path
from docx import Document
from .base_parser import BaseParser, ResumeData, Experience, Education, Certification, Project


class DOCXParser(BaseParser):
    """Parser for DOCX resume files"""
    
    SUPPORTED_FORMATS = ["docx"]
    
    def parse(self) -> ResumeData:
        """Parse DOCX resume and extract all relevant information"""
        self.raw_text = self._extract_text()
        resume_data = ResumeData()
        
        resume_data.raw_text = self.raw_text
        resume_data.file_path = self.file_path
        resume_data.file_type = "docx"
        
        # Extract basic information
        resume_data.email = self._extract_email(self.raw_text)
        resume_data.phone = self._extract_phone(self.raw_text)
        links = self._extract_links(self.raw_text)
        resume_data.linkedin = links.get("linkedin")
        resume_data.github = links.get("github")
        resume_data.website = links.get("website")
        
        # Extract name
        resume_data.full_name = self._extract_name(self.raw_text)
        
        # Extract sections
        resume_data.summary = self._extract_summary(self.raw_text)
        resume_data.technical_skills = self._extract_skills(self.raw_text)
        resume_data.experiences = self._extract_experiences(self.raw_text)
        resume_data.education = self._extract_education(self.raw_text)
        resume_data.certifications = self._extract_certifications(self.raw_text)
        resume_data.projects = self._extract_projects(self.raw_text)
        
        return resume_data
    
    def _extract_text(self) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(str(self.file_path))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return self._clean_text(" ".join(text_parts))
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from text"""
        lines = text.split('\n')
        if lines:
            for i, line in enumerate(lines[:10]):
                line = line.strip()
                if line and len(line.split()) <= 4:
                    if '@' not in line and not any(c.isdigit() for c in line):
                        # Skip lines that look like section headers
                        skip_words = ['summary', 'experience', 'education', 'skills', 
                                     'certifications', 'projects', 'contact', 'email',
                                     'phone', 'linkedin', 'github']
                        if line.lower() not in skip_words:
                            return line
        return ""
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary"""
        import re
        summary_patterns = [
            r'(?i)(professional\s*summary|summary|profile|objective|about)[:\s]*(.+?)(?=\n\n|\nexperience|\neducation|\nskills|\ncertifications|\nprojects|$)',
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                summary = match.group(2).strip()
                summary = re.sub(r'\s+', ' ', summary)
                return summary[:500]
        return ""
    
    def _extract_skills(self, text: str) -> list:
        """Extract skills from resume text"""
        from config.config import SKILL_CATEGORIES
        
        all_known_skills = []
        for category in SKILL_CATEGORIES.values():
            all_known_skills.extend([s.lower() for s in category])
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in all_known_skills:
            if skill in text_lower:
                for category_skills in SKILL_CATEGORIES.values():
                    for original_skill in category_skills:
                        if original_skill.lower() == skill and original_skill not in found_skills:
                            found_skills.append(original_skill)
                            break
        
        return list(set(found_skills))
    
    def _extract_section(self, text: str, section_name: str) -> str:
        """Extract a specific section from the resume"""
        import re
        section_pattern = rf'(?i)(?:^|\n)\s*{section_name}\s*[:\-]?\s*(.+?)(?=\n\n|\n(?:[A-Z]|\d|\*))'
        match = re.search(section_pattern, text, re.DOTALL)
        if match:
            return match.group(1).strip()
        return ""
    
    def _extract_experiences(self, text: str) -> list:
        """Extract work experience entries"""
        experiences = []
        
        exp_section = self._extract_section(text, 'experience')
        
        if exp_section:
            import re
            
            # Pattern for date ranges
            date_pattern = (
                r'(?i)(\\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\\.?\\s*\\d{2,4}|'
                r'\\d{1,2}/\\d{2,4}|\\d{4})\\s*[-–—to]+\\s*'
                r'(\\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\\.?\\s*\\d{2,4}|'
                r'\\d{1,2}/\\d{2,4}|\\d{4}|present|current|now)'
            )
            
            entries = re.split(date_pattern, exp_section)
            
            for i in range(0, len(entries), 3):
                if i + 2 < len(entries):
                    duration = f"{entries[i].strip()} - {entries[i+1].strip()}"
                    description = entries[i+2].strip()
                    
                    lines = description.split('\n')
                    first_line = lines[0] if lines else ""
                    role_company = self._parse_role_company(first_line)
                    
                    exp = Experience(
                        role=role_company.get("role", ""),
                        company=role_company.get("company", ""),
                        duration=duration,
                        description=description[:500]
                    )
                    experiences.append(exp)
        
        return experiences
    
    def _parse_role_company(self, text: str) -> dict:
        """Parse role and company from text line"""
        import re
        
        patterns = [
            r'(.+?)\s+(?:at|@|with|for)\s+(.+)',
            r'(.+?)\s+[-–—]\s+(.+)',
        ]
        
        result = {"role": "", "company": ""}
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result["role"] = match.group(1).strip()
                result["company"] = match.group(2).strip()
                break
        
        if not result["role"]:
            result["role"] = text.strip()
        
        return result
    
    def _extract_education(self, text: str) -> list:
        """Extract education entries"""
        education = []
        
        edu_section = self._extract_section(text, 'education')
        
        if edu_section:
            import re
            
            degree_patterns = [
                r'(?i)(bachelor['']?s?|master['']?s?|ph\.?d\.?|mba|b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.?|associate['']?s?)\s+(?:of|in)?\s*(.+?)(?=\n|$)',
            ]
            
            lines = edu_section.split('\n')
            current_entry = {}
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                for pattern in degree_patterns:
                    match = re.search(pattern, line)
                    if match:
                        degree = match.group(1) if match.group(1) else match.group(0)
                        institution = match.group(2) if match.group(2) else ""
                        current_entry["degree"] = degree.title()
                        current_entry["institution"] = institution.title()
                        break
                
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match:
                    current_entry["year"] = year_match.group()
                
                if current_entry.get("degree") and current_entry.get("institution"):
                    edu = Education(
                        degree=current_entry.get("degree", ""),
                        institution=current_entry.get("institution", ""),
                        year=current_entry.get("year", "")
                    )
                    education.append(edu)
                    current_entry = {}
        
        return education
    
    def _extract_certifications(self, text: str) -> list:
        """Extract certification entries"""
        certifications = []
        
        cert_section = self._extract_section(text, 'certifications')
        
        if cert_section:
            lines = cert_section.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                import re
                issuer_match = re.search(r'(?i)(aws|azure|google|microsoft|cisco|oracle|ibm|itil|pmp|pmo)(?:\s|$)', line)
                issuer = issuer_match.group(1).title() if issuer_match else ""
                
                date_match = re.search(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}\b', line, re.IGNORECASE)
                date = date_match.group() if date_match else ""
                
                cert = Certification(
                    name=line.split(issuer)[0].strip() if issuer else line,
                    issuer=issuer,
                    date=date
                )
                certifications.append(cert)
        
        return certifications
    
    def _extract_projects(self, text: str) -> list:
        """Extract project entries"""
        projects = []
        
        project_section = self._extract_section(text, 'projects')
        
        if project_section:
            import re
            lines = project_section.split('\n')
            current_project = {}
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if not current_project.get("name"):
                    if not line.startswith(('•', '-', '*', '1.', '2.', '3.')):
                        current_project["name"] = line
                        continue
                
                tech_pattern = r'(?i)(python|java|javascript|react|node|django|flask|tensorflow|pytorch|aws|docker|kubernetes|mongodb|postgresql|sql)'
                techs = re.findall(tech_pattern, line)
                if techs:
                    current_project["technologies"] = [t.title() for t in techs]
                
                if "description" not in current_project:
                    current_project["description"] = ""
                current_project["description"] += line + " "
                
                if current_project.get("name") and len(current_project.get("description", "")) > 20:
                    proj = Project(
                        name=current_project.get("name", ""),
                        description=current_project.get("description", "").strip(),
                        technologies=current_project.get("technologies", [])
                    )
                    projects.append(proj)
                    current_project = {}
        
        return projects
